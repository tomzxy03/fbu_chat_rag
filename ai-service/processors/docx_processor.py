"""DOCX processor – sử dụng python-docx."""

from docx import Document
from .base import BaseProcessor


class DocxProcessor(BaseProcessor):
    SUPPORTED_EXTENSIONS = [".docx", ".doc"]

    def extract_text(self, file_path: str) -> str:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Trích xuất từ tables (thường có trong quy chế trường)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n".join(paragraphs)
